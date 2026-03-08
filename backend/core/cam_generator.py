"""
Five Cs CAM Report Generator with PDF/DOCX export.
Generates Credit Appraisal Memo structured around the Five Cs of Credit:
Character, Capacity, Capital, Collateral, Conditions.

Professional fallback messaging for missing data.
Includes Data Quality Summary section.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

# ─── Professional Fallback Constants ─────────────────────────────────────────

_UNAVAILABLE = "Data not available from uploaded documents"
_NO_DATA = "No data available"


def _fallback(cp: dict, default_msg: str = _UNAVAILABLE) -> str:
    """Professional fallback for a checkpoint result — never prints N/A."""
    label = cp.get("result_label", "")
    if not label or label == "N/A":
        return default_msg
    if cp.get("data_missing"):
        return default_msg
    return label


def _fallback_tier(cp: dict) -> str:
    """Professional fallback for score tier."""
    tier = cp.get("score_tier", "")
    if not tier or tier == "N/A":
        return "Not assessed"
    if cp.get("data_missing"):
        return "Not assessed (data unavailable)"
    return tier


# ─── Five Cs CAM Builder ─────────────────────────────────────────────────────


def build_cam_report(
    session: dict,
    forensic_results: dict,
    extracted_data: dict,
) -> dict:
    """
    Build a structured Five Cs CAM report from session data and forensic results.
    Returns a dict with sections for each C plus decision and full narrative.
    """
    fd = session.get("financial_data", {})
    company = session.get("company_name", "Unknown Company")
    loan_amount = session.get("loan_amount", 0)
    loan_purpose = session.get("loan_purpose", "Business expansion")

    results = forensic_results.get("results", [])
    aggregate = forensic_results.get("aggregate_score", 0)
    risk_grade = forensic_results.get("risk_grade", "Pending")
    vetoed = forensic_results.get("vetoed", False)
    data_meta = extracted_data.get("_meta", {})

    # Helper to find checkpoint result by ID
    def cp_result(cp_id: int) -> dict:
        for r in results:
            if r["id"] == cp_id:
                return r
        return {}

    gst = extracted_data.get("GST", {})
    itr = extracted_data.get("ITR_Financials", {})
    bank = extracted_data.get("Banking", {})
    ml = extracted_data.get("Management_Legal", {})

    # ── 1. CHARACTER ──────────────────────────────────────────────────────
    character_items = []
    cp1 = cp_result(1)
    cp2 = cp_result(2)
    cp7 = cp_result(7)
    cp10 = cp_result(10)
    cp19 = cp_result(19)
    cp25 = cp_result(25)

    character_items.append(f"GST Filing Discipline: {_fallback(cp1, 'Filing data not available')} ({_fallback_tier(cp1)})")
    character_items.append(f"Sales Integrity Variance: {_fallback(cp2, 'Sales comparison data not available')} ({_fallback_tier(cp2)})")
    character_items.append(f"Revenue Integrity (ITR vs GST): {_fallback(cp7, 'Revenue cross-verification data unavailable')} ({_fallback_tier(cp7)})")
    character_items.append(f"Tax Compliance: {_fallback(cp10, 'Tax compliance data not available')} ({_fallback_tier(cp10)})")
    character_items.append(f"Auditor Opinion: {_fallback(cp19, 'Auditor report not provided')} ({_fallback_tier(cp19)})")
    character_items.append(f"Litigation Exposure: {_fallback(cp25, 'No litigation data available')} ({_fallback_tier(cp25)})")

    char_scores = [cp1, cp2, cp7, cp10, cp19, cp25]
    char_avg = _avg_tier(char_scores)
    trust_label = "HIGH" if char_avg >= 8 else ("MODERATE" if char_avg >= 5 else "LOW")

    character = {
        "assessment": character_items,
        "trust_score": trust_label,
        "summary": f"Character Assessment: Trust Score = {trust_label}",
    }

    # ── 2. CAPACITY ──────────────────────────────────────────────────────
    cp9 = cp_result(9)
    dscr_val = cp9.get("result_value")
    np_ = itr.get("net_profit")
    dep = itr.get("depreciation")
    intr = itr.get("interest_expense")
    prin = itr.get("principal_repayment")

    loan_display = _format_inr(loan_amount)

    # Build DSCR display
    if dscr_val is None or dscr_val == "Data Missing" or dscr_val == "Not Computable":
        missing_parts = []
        if prin is None:
            missing_parts.append("principal repayment")
        if np_ is None:
            missing_parts.append("net profit")
        if dep is None:
            missing_parts.append("depreciation")
        if intr is None:
            missing_parts.append("interest expense")
        if missing_parts:
            dscr_display = f"DSCR: Not computable ({', '.join(missing_parts)} data unavailable)"
        else:
            dscr_display = "DSCR: Not computable (insufficient data)"
        dscr_num = None
    else:
        dscr_num = dscr_val if isinstance(dscr_val, (int, float)) else None
        dscr_display = f"DSCR = {dscr_val}"

    capacity_lines = [dscr_display]
    if np_ is not None:
        capacity_lines.append(f"Net Profit: {_format_inr(np_)}")
    if dep is not None:
        capacity_lines.append(f"Depreciation: {_format_inr(dep)}")
    if intr is not None:
        capacity_lines.append(f"Interest: {_format_inr(intr)}")

    if dscr_num is not None and dscr_num >= 1.5:
        capacity_lines.append(f"The company demonstrates strong repayment ability and can comfortably service a {loan_display} loan.")
    elif dscr_num is not None and dscr_num >= 1.2:
        capacity_lines.append(f"The company has adequate repayment capacity for a {loan_display} loan with moderate buffer.")
    elif dscr_num is not None:
        capacity_lines.append(f"Repayment capacity is strained for a {loan_display} loan — additional safeguards recommended.")
    else:
        capacity_lines.append(f"Repayment capacity assessment requires principal repayment and interest data.")

    capacity = {
        "dscr": dscr_val if dscr_val is not None else "Not Computable",
        "net_profit": np_,
        "depreciation": dep,
        "interest_expense": intr,
        "summary": "\n".join(capacity_lines),
    }

    # ── 3. CAPITAL ───────────────────────────────────────────────────────
    debt = itr.get("total_liabilities")
    equity = itr.get("net_worth")

    if debt is not None and equity is not None and equity > 0:
        de_ratio = round(debt / equity, 2)
        capital_summary = (
            f"Debt-Equity Ratio = {de_ratio}\n"
            f"Total Liabilities: {_format_inr(debt)} | Net Worth: {_format_inr(equity)}\n"
        )
        if de_ratio <= 2:
            capital_summary += "Leverage is low and financial solvency is strong."
        elif de_ratio <= 4:
            capital_summary += "Moderate leverage — monitor debt servicing closely."
        else:
            capital_summary += "High leverage — significant solvency risk."
    else:
        de_ratio = None
        capital_summary = "Debt-equity analysis not computable — liability or net worth data unavailable."

    capital = {
        "debt_equity_ratio": de_ratio if de_ratio is not None else "Not Computable",
        "total_liabilities": debt,
        "net_worth": equity,
        "summary": capital_summary,
    }

    # ── 4. COLLATERAL ────────────────────────────────────────────────────
    cp23 = cp_result(23)
    cp29 = cp_result(29)
    fixed_assets = itr.get("fixed_assets")

    if fixed_assets is not None and loan_amount and loan_amount > 0:
        coverage = round((fixed_assets / loan_amount) * 100, 1)
    else:
        coverage = None

    collateral = {
        "primary_asset_coverage": _format_inr(fixed_assets) if fixed_assets is not None else _UNAVAILABLE,
        "collateral_coverage_pct": coverage if coverage is not None else "Not Computable",
        "share_pledge": _fallback(cp23, "Share pledge data not available"),
        "lien_status": _fallback(cp29, "Lien verification data not available"),
        "summary": (
            f"Primary asset coverage = {_format_inr(fixed_assets) if fixed_assets is not None else _UNAVAILABLE}\n"
            f"Collateral coverage = {f'{coverage}%' if coverage is not None else 'Not computable (fixed asset or loan amount data unavailable)'}\n"
            f"Share pledge: {_fallback(cp23, 'Share pledge data not available')}\n"
            f"Lien status: {_fallback(cp29, 'Lien verification data not available')}"
        ),
    }

    # ── 5. CONDITIONS ────────────────────────────────────────────────────
    cp6 = cp_result(6)
    cp27 = cp_result(27)
    cp28 = cp_result(28)
    cp3 = cp_result(3)

    conditions = {
        "business_age": _fallback(cp6, "Business age data not available"),
        "operational_health": _fallback(cp27, "Operational data not available"),
        "news_sentiment": _fallback(cp28, "Market sentiment data not available"),
        "turnover_trend": _fallback(cp3, "Turnover trend data not available"),
        "summary": (
            f"Business Age: {_fallback(cp6, 'Business age data not available')}\n"
            f"Operational Health: {_fallback(cp27, 'Operational data not available')}\n"
            f"Market Sentiment: {_fallback(cp28, 'Market sentiment data not available')}\n"
            f"Turnover Trend: {_fallback(cp3, 'Turnover trend data not available')}\n"
            f"Market conditions are {'favorable' if cp28.get('score_tier') == 'Score_3' and not cp28.get('data_missing') else 'not assessed due to limited data'}."
        ),
    }

    # ── DECISION ─────────────────────────────────────────────────────────
    if vetoed:
        decision_status = "REJECTED"
        veto_cp = forensic_results.get("veto_checkpoint", {})
        decision_reason = (
            f"HARD VETO triggered at Checkpoint {veto_cp.get('id', '?')} ({veto_cp.get('name', '?')}): "
            f"{_fallback(veto_cp, 'Critical threshold breached')}"
        )
        rec_rate = "Not applicable"
        rec_tenure = "Not applicable"
        rec_limit = "Not applicable"
    elif risk_grade == "Prime Borrower":
        decision_status = "APPROVED"
        rec_rate = "10.00%"
        rec_tenure = "7 Years"
        rec_limit = loan_display
        decision_reason = f"Strong DSCR ({dscr_val}), clean compliance, healthy liquidity, low litigation exposure."
    elif risk_grade == "Strong Borrower":
        decision_status = "APPROVED"
        rec_rate = "10.75%"
        rec_tenure = "5 Years"
        rec_limit = loan_display
        decision_reason = f"Good DSCR ({dscr_val}), acceptable compliance, adequate liquidity."
    elif risk_grade == "Moderate Risk":
        decision_status = "CONDITIONAL APPROVAL"
        rec_rate = "12.50%"
        rec_tenure = "3 Years"
        rec_limit = _format_inr(loan_amount * 0.75) if loan_amount else "Not applicable"
        decision_reason = f"Moderate risk detected. DSCR = {dscr_val or 'Not Computable'}. Enhanced monitoring recommended."
    else:
        decision_status = "REJECTED"
        rec_rate = "Not applicable"
        rec_tenure = "Not applicable"
        rec_limit = "Not applicable"
        decision_reason = f"Risk score {aggregate}/300 is below the acceptance threshold of 180."

    decision = {
        "status": decision_status,
        "loan_limit": rec_limit,
        "interest_rate": rec_rate,
        "tenure": rec_tenure,
        "risk_grade": risk_grade,
        "aggregate_score": aggregate,
        "reason": decision_reason,
    }

    # ── DATA QUALITY SUMMARY ─────────────────────────────────────────────
    data_quality = {
        "present_count": data_meta.get("present_count", 0),
        "missing_count": data_meta.get("missing_count", 0),
        "total_required": data_meta.get("total_required", 0),
        "data_completeness_pct": data_meta.get("data_completeness_pct", 0),
        "missing_fields": data_meta.get("missing_fields", []),
    }

    # ── FIVE CS SUMMARY ──────────────────────────────────────────────────
    five_cs_summary = {
        "Character": trust_label,
        "Capacity": _fallback_tier(cp9) if cp9 else "Not assessed",
        "Capital": ("Strong" if de_ratio is not None and de_ratio <= 2
                    else ("Moderate" if de_ratio is not None and de_ratio <= 4
                          else ("Weak" if de_ratio is not None else "Not assessed"))),
        "Collateral": f"{coverage}%" if coverage is not None else "Not assessed",
        "Conditions": _fallback_tier(cp28) if cp28 else "Not assessed",
    }

    # ── FULL NARRATIVE ───────────────────────────────────────────────────
    narrative = _build_narrative(
        company, character, capacity, capital, collateral, conditions, decision, data_quality
    )

    return {
        "company_name": company,
        "loan_amount": loan_amount,
        "loan_purpose": loan_purpose,
        "generated_at": datetime.utcnow().isoformat(),
        "character": character,
        "capacity": capacity,
        "capital": capital,
        "collateral": collateral,
        "conditions": conditions,
        "decision": decision,
        "data_quality": data_quality,
        "five_cs_summary": five_cs_summary,
        "full_narrative": narrative,
    }


def _avg_tier(checkpoints: list[dict]) -> float:
    """Average score points across checkpoints."""
    pts = [cp.get("score_points", 0) for cp in checkpoints if cp]
    return sum(pts) / len(pts) if pts else 0


def _format_inr(amount: float | int | None) -> str:
    """Format amount in Indian convention."""
    if amount is None:
        return "Data not available"
    if not amount:
        return "₹0"
    try:
        amt = float(amount)
    except (TypeError, ValueError):
        return str(amount)
    if amt >= 10000000:
        return f"₹{amt / 10000000:.2f} Crore"
    elif amt >= 100000:
        return f"₹{amt / 100000:.2f} Lakhs"
    else:
        return f"₹{amt:,.0f}"


def _build_narrative(company, character, capacity, capital, collateral, conditions, decision, data_quality) -> str:
    """Build a full prose CAM narrative with Data Quality Summary."""
    sections = []
    sections.append(f"CREDIT APPRAISAL MEMORANDUM — {company}")
    sections.append(f"Generated: {datetime.utcnow().strftime('%d %B %Y %H:%M UTC')}")
    sections.append("=" * 60)

    # Data Quality Header
    dq = data_quality
    comp_pct = dq.get("data_completeness_pct", 0)
    sections.append(f"\nDATA COMPLETENESS: {comp_pct}%")
    sections.append(f"Extracted Fields: {dq.get('present_count', 0)} / {dq.get('total_required', 0)}")
    sections.append(f"Missing Fields: {dq.get('missing_count', 0)}")
    if dq.get("missing_fields"):
        sections.append("\nMissing Variables:")
        for mf in dq["missing_fields"][:20]:
            sections.append(f"  - {mf}")

    sections.append("\n" + "=" * 60)

    sections.append("\n1. CHARACTER ASSESSMENT")
    sections.append("-" * 40)
    for item in character["assessment"]:
        sections.append(f"  • {item}")
    sections.append(f"\n  Trust Score: {character['trust_score']}")

    sections.append("\n2. CAPACITY (Repayment Ability)")
    sections.append("-" * 40)
    sections.append(f"  {capacity['summary']}")

    sections.append("\n3. CAPITAL (Financial Leverage)")
    sections.append("-" * 40)
    sections.append(f"  {capital['summary']}")

    sections.append("\n4. COLLATERAL")
    sections.append("-" * 40)
    sections.append(f"  {collateral['summary']}")

    sections.append("\n5. CONDITIONS (Market & External)")
    sections.append("-" * 40)
    sections.append(f"  {conditions['summary']}")

    sections.append("\n" + "=" * 60)
    sections.append("DATA QUALITY SUMMARY")
    sections.append("-" * 40)
    sections.append(f"  Extracted Financial Fields: {dq.get('present_count', 0)} / {dq.get('total_required', 0)}")
    sections.append(f"  Data Completeness: {comp_pct}%")
    if dq.get("missing_fields"):
        sections.append(f"  Missing ({dq.get('missing_count', 0)} variables):")
        for mf in dq["missing_fields"][:20]:
            sections.append(f"    - {mf}")

    sections.append("\n" + "=" * 60)
    sections.append(f"DECISION: {decision['status']}")
    sections.append("-" * 40)
    sections.append(f"  Loan Limit: {decision['loan_limit']}")
    sections.append(f"  Interest Rate: {decision['interest_rate']}")
    sections.append(f"  Tenure: {decision['tenure']}")
    sections.append(f"  Risk Grade: {decision['risk_grade']} ({decision['aggregate_score']}/300)")
    sections.append(f"\n  Reason: {decision['reason']}")

    return "\n".join(sections)


# ─── PDF Generator ───────────────────────────────────────────────────────────


def generate_pdf(cam_data: dict) -> bytes:
    """Generate a PDF version of the CAM report using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    )
    from reportlab.lib.units import inch, mm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=30*mm, bottomMargin=20*mm,
                            leftMargin=20*mm, rightMargin=20*mm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CamTitle", parent=styles["Title"], fontSize=18,
                                  textColor=HexColor("#0F172A"), spaceAfter=6)
    heading_style = ParagraphStyle("CamHeading", parent=styles["Heading2"], fontSize=13,
                                    textColor=HexColor("#0F172A"), spaceBefore=16, spaceAfter=6)
    body_style = ParagraphStyle("CamBody", parent=styles["Normal"], fontSize=10,
                                 leading=14, textColor=HexColor("#334155"))
    label_style = ParagraphStyle("CamLabel", parent=styles["Normal"], fontSize=9,
                                  textColor=HexColor("#64748B"), spaceBefore=2)

    story = []

    # Title
    company = cam_data.get("company_name", "Company")
    story.append(Paragraph(f"Credit Appraisal Memorandum", title_style))
    story.append(Paragraph(f"{company}", body_style))
    story.append(Paragraph(f"Generated: {cam_data.get('generated_at', '')[:10]}", label_style))

    # Data Quality badge
    dq = cam_data.get("data_quality", {})
    comp_pct = dq.get("data_completeness_pct", 0)
    story.append(Paragraph(f"Data Completeness: {comp_pct}% ({dq.get('present_count', 0)}/{dq.get('total_required', 0)} fields)", label_style))
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", color=HexColor("#E2E8F0")))

    # Decision summary
    decision = cam_data.get("decision", {})
    story.append(Heading(heading_style, "Decision"))
    dec_data = [
        ["Status", decision.get("status", "")],
        ["Loan Limit", str(decision.get("loan_limit", ""))],
        ["Interest Rate", str(decision.get("interest_rate", ""))],
        ["Tenure", str(decision.get("tenure", ""))],
        ["Risk Grade", f"{decision.get('risk_grade', '')} ({decision.get('aggregate_score', 0)}/300)"],
    ]
    story.append(_make_table(dec_data))
    story.append(Paragraph(f"<b>Reason:</b> {decision.get('reason', '')}", body_style))
    story.append(Spacer(1, 8))

    # Five Cs sections
    for section_key, section_title in [("character", "1. Character"), ("capacity", "2. Capacity"),
                                        ("capital", "3. Capital"), ("collateral", "4. Collateral"),
                                        ("conditions", "5. Conditions")]:
        section = cam_data.get(section_key, {})
        story.append(Paragraph(section_title, heading_style))
        summary = section.get("summary", "")
        for line in summary.split("\n"):
            if line.strip():
                story.append(Paragraph(line.strip(), body_style))
        story.append(Spacer(1, 4))

    # Data Quality Summary
    story.append(Paragraph("Data Quality Summary", heading_style))
    dq_data = [
        ["Extracted Fields", f"{dq.get('present_count', 0)} / {dq.get('total_required', 0)}"],
        ["Missing Fields", str(dq.get('missing_count', 0))],
        ["Completeness", f"{comp_pct}%"],
    ]
    story.append(_make_table(dq_data))
    if dq.get("missing_fields"):
        missing_text = ", ".join(dq["missing_fields"][:15])
        story.append(Paragraph(f"<b>Missing:</b> {missing_text}", body_style))

    doc.build(story)
    return buf.getvalue()


def Heading(style, text):
    """Helper to create a heading Paragraph."""
    from reportlab.platypus import Paragraph
    return Paragraph(text, style)


def _make_table(data):
    """Create a simple two-column table."""
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch

    t = Table(data, colWidths=[1.8*inch, 3.5*inch])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), HexColor("#64748B")),
        ("TEXTCOLOR", (1, 0), (1, -1), HexColor("#0F172A")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("LINEBELOW", (0, -1), (-1, -1), 0.5, HexColor("#E2E8F0")),
    ]))
    return t


# ─── DOCX Generator ─────────────────────────────────────────────────────────


def generate_docx(cam_data: dict) -> bytes:
    """Generate a DOCX version of the CAM report using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    company = cam_data.get("company_name", "Company")
    title = doc.add_heading("Credit Appraisal Memorandum", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(company)
    doc.add_paragraph(f"Generated: {cam_data.get('generated_at', '')[:10]}")

    # Data Quality badge
    dq = cam_data.get("data_quality", {})
    comp_pct = dq.get("data_completeness_pct", 0)
    doc.add_paragraph(f"Data Completeness: {comp_pct}% ({dq.get('present_count', 0)}/{dq.get('total_required', 0)} fields)")
    doc.add_paragraph("")

    # Decision
    doc.add_heading("Decision", level=1)
    decision = cam_data.get("decision", {})
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light List"
    rows_data = [
        ("Status", decision.get("status", "")),
        ("Loan Limit", str(decision.get("loan_limit", ""))),
        ("Interest Rate", str(decision.get("interest_rate", ""))),
        ("Tenure", str(decision.get("tenure", ""))),
        ("Risk Grade", f"{decision.get('risk_grade', '')} ({decision.get('aggregate_score', 0)}/300)"),
    ]
    for i, (label, value) in enumerate(rows_data):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value)

    doc.add_paragraph(f"\nReason: {decision.get('reason', '')}")

    # Five Cs
    for section_key, section_title in [("character", "1. Character"), ("capacity", "2. Capacity"),
                                        ("capital", "3. Capital"), ("collateral", "4. Collateral"),
                                        ("conditions", "5. Conditions")]:
        doc.add_heading(section_title, level=1)
        section = cam_data.get(section_key, {})
        summary = section.get("summary", "")
        for line in summary.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    # Data Quality Summary
    doc.add_heading("Data Quality Summary", level=1)
    doc.add_paragraph(f"Extracted Financial Fields: {dq.get('present_count', 0)} / {dq.get('total_required', 0)}")
    doc.add_paragraph(f"Data Completeness: {comp_pct}%")
    doc.add_paragraph(f"Missing Fields: {dq.get('missing_count', 0)}")
    if dq.get("missing_fields"):
        for mf in dq["missing_fields"][:15]:
            doc.add_paragraph(f"• {mf}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
